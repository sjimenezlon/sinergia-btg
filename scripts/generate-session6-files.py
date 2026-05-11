"""
Genera los archivos descargables de la Sesión 6 — Quick Wins.
Datos sintéticos y verosímiles para que cada ejercicio tenga insumo directo.
Ningún dato real de BTG ni de clientes.
"""
import os
import random
from datetime import datetime, timedelta

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(42)

OUT = os.path.join(os.path.dirname(__file__), "..", "public", "sesion-6")
os.makedirs(OUT, exist_ok=True)


# ──────────────────────────────────────────────────────────────────────────
# Estilos compartidos
# ──────────────────────────────────────────────────────────────────────────
HEADER_FILL = PatternFill("solid", fgColor="0D1229")
HEADER_FONT = Font(bold=True, color="00D4E5", size=11)
SUB_FILL = PatternFill("solid", fgColor="151A3A")
THIN = Side(border_style="thin", color="2A2F4F")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_header(ws, row, cols):
    for col in range(1, cols + 1):
        c = ws.cell(row=row, column=col)
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = BORDER


def autosize(ws):
    from openpyxl.cell.cell import MergedCell
    widths = {}
    for row in ws.iter_rows():
        for c in row:
            if isinstance(c, MergedCell) or c.value is None:
                continue
            widths[c.column_letter] = max(widths.get(c.column_letter, 10), len(str(c.value)) + 2)
    for letter, w in widths.items():
        ws.column_dimensions[letter].width = min(w, 60)


def add_heading(doc, text, level=1, color="00D4E5"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def banner(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("E85A1F")


# ──────────────────────────────────────────────────────────────────────────
# 1. Utilities Colombia 2026 — research brief para NotebookLM
# ──────────────────────────────────────────────────────────────────────────
def make_utilities_brief():
    doc = Document()
    add_heading(doc, "Utilities Colombia 2026 · Research Brief", level=0)
    banner(doc, "DATOS SINTÉTICOS — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")
    add_para(doc,
        "Documento consolidado de cinco fuentes equivalentes (reportes trimestrales de ISA, EPM, "
        "Celsia, informe sectorial de research bursátil y transcripción de panel sectorial). "
        "Pensado para cargarse en NotebookLM como base única, simulando un dataroom corto.",
        italic=True, size=10)

    add_heading(doc, "1. Panorama sectorial (síntesis)", level=1)
    add_para(doc,
        "El sector utilities colombiano cerró 2025 con un Ebitda agregado de COP 18,4 billones "
        "(+9,1% a/a) y un capex agregado de COP 11,2 billones, concentrado en transmisión (45%) y "
        "renovables (32%). La inflación promedio anual del IPP energético fue de 6,8%, ligeramente "
        "por encima del 5,9% del IPC, con un efecto neto positivo sobre tarifas indexadas. La tasa "
        "BanRep cerró en 9,25% (TPM) y el consenso de research la proyecta en 7,25% a diciembre 2026.")

    add_heading(doc, "2. Emisores cubiertos", level=1)
    emisores = [
        ("ISA", "Transmisión LatAm",
         "Backbone regional. Capex 2025: COP 4,1 B (Brasil 52%, Colombia 28%, Perú/Chile 20%). "
         "EBITDA margen 71%. Cobertura intereses 5,2x. Riesgo regulatorio bajo, FX hedging 78% deuda. "
         "Próximo bullet 2027: USD 500 M senior unsecured."),
        ("EPM", "Multi-utility integrado",
         "Hidroituango sumando 1.200 MW al SIN. Subsidios cruzados generación–distribución bajo "
         "revisión por CREG. Deuda/EBITDA 3,8x (objetivo ≤3,5x a 2027). Calificación BB+ S&P (perspectiva "
         "estable). Riesgo: presión política sobre tarifas en estratos 1-3."),
        ("Celsia", "Generación + distribución",
         "Mix 67% renovable. Granjas solares Yumbo, Bolívar, Tolima. PPA promedio 8,2 años. "
         "Margen EBITDA 41%. Próxima emisión local COP 600 mil M para refi 2027."),
        ("Promigas", "Gas natural · transporte y distribución",
         "Caja recurrente alta, pero exposición a declinación de reservas (Guajira). Inversión en "
         "regasificación Buenaventura. Cobertura intereses 4,1x. Rating BBB– Fitch."),
        ("Enel Colombia", "Generación + comercialización",
         "Foco en respaldo térmico y renovables nuevas. Plan capex 2026-30: USD 2,4 B. Sensibilidad "
         "alta a precios spot durante El Niño."),
    ]
    for nombre, sub, txt in emisores:
        add_para(doc, f"{nombre} — {sub}", bold=True, size=11, color="E85A1F")
        add_para(doc, txt)

    add_heading(doc, "3. Riesgos regulatorios 2026", level=1)
    for i, t in enumerate([
        "Revisión tarifaria CREG Resolución 015/2018 — recalibración WACC regulado, posible recorte 60-90 bps.",
        "Aplazamiento subasta cargo por confiabilidad — afecta valoración de plantas térmicas marginales.",
        "Discusión congresal Ley 142 — riesgo de topes a contribución de solidaridad estratos 5-6.",
        "PND 2026-30 — meta 25% renovables no convencionales sobre matriz, exige fast-track ambiental.",
        "Tarifa de transmisión STN — ajuste por demanda real 2024-25 podría revisar ingresos regulados de ISA Intercolombia.",
    ], start=1):
        add_para(doc, f"R{i}. {t}")

    add_heading(doc, "4. Indicadores clave (cierre 2025)", level=1)
    add_para(doc,
        "Demanda SIN: 80,4 TWh (+3,1% a/a). Generación hidráulica 64,2%; térmica 18,1%; renovables "
        "no convencionales 12,8%; otras 4,9%. Aporte hídrico al embalse: 102% del promedio histórico. "
        "Tarifa promedio CU regulado: COP 869/kWh (+8,2% a/a). Pérdidas técnicas + no técnicas: 9,4%.",)

    add_heading(doc, "5. Vista equipo Renta Fija — qué mirar", level=1)
    for t in [
        "Comparar D/EBITDA ISA vs EPM vs Celsia para asignar peso relativo en portafolio investment grade.",
        "Monitorear emisión Celsia COP 600 mil M (mayo 2026): cupón y demanda son referencia para el mercado local.",
        "Hidroituango: revisar covenants de la sindicada EPM si flujo libre cae por debajo del threshold 1,1x.",
        "Riesgo regulatorio CREG — pricing del bono ISA 2032 podría requerir spread adicional 15-25 bps en escenario adverso.",
        "Curva CETIP equivalente (TES UVR) — mejor pricing referencial que LIBOR para emisores hidro indexados a IPC.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10)

    add_heading(doc, "6. Glosario de términos", level=1)
    for k, v in [
        ("CREG", "Comisión de Regulación de Energía y Gas (Colombia)."),
        ("SIN", "Sistema Interconectado Nacional."),
        ("STN", "Sistema de Transmisión Nacional."),
        ("Cargo por Confiabilidad", "Subasta de respaldo firme administrada por CREG."),
        ("WACC regulado", "Costo de capital reconocido en tarifa por la CREG."),
    ]:
        add_para(doc, f"{k} — {v}")

    doc.save(os.path.join(OUT, "01-utilities-colombia-brief.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 2. Portafolio abril 2026 — 120 posiciones (Excel)
# ──────────────────────────────────────────────────────────────────────────
def make_portafolio_abril():
    wb = Workbook()
    ws = wb.active
    ws.title = "Posiciones"

    ws["A1"] = "BTG · Portafolio Renta Fija LatAm · Abril 2026 (DATOS SINTÉTICOS)"
    ws["A1"].font = Font(bold=True, color="E85A1F", size=12)
    ws.merge_cells("A1:J1")

    headers = ["ticker", "emisor", "sector", "pais", "rating", "duration_años",
               "ytm_%", "nominal_mm_cop", "precio_clean", "valor_mercado_mm_cop"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=3, column=i, value=h)
    style_header(ws, 3, len(headers))

    emisores_pool = [
        ("ISA", "Utilities", "CO", "AAA"),
        ("EPM", "Utilities", "CO", "AA+"),
        ("CELSIA", "Utilities", "CO", "AA"),
        ("PROMIGAS", "Energía", "CO", "AA"),
        ("ECOPETROL", "Energía", "CO", "BBB"),
        ("BANCOLOMBIA", "Bancos", "CO", "BBB-"),
        ("DAVIVIENDA", "Bancos", "CO", "BB+"),
        ("BBVA COL", "Bancos", "CO", "BBB-"),
        ("GRUPO ARGOS", "Holding", "CO", "AA-"),
        ("CEMARGOS", "Cementos", "CO", "BBB"),
        ("NUTRESA", "Consumo", "CO", "BBB+"),
        ("AVAL", "Holding", "CO", "BB+"),
        ("PETROBRAS", "Energía", "BR", "BBB-"),
        ("VALE", "Minería", "BR", "BBB"),
        ("ITAU", "Bancos", "BR", "BBB-"),
        ("CEMEX", "Cementos", "MX", "BB+"),
        ("FEMSA", "Consumo", "MX", "A-"),
        ("BIMBO", "Consumo", "MX", "BBB+"),
        ("PEMEX", "Energía", "MX", "BB-"),
        ("CODELCO", "Minería", "CL", "A"),
        ("BCI", "Bancos", "CL", "A-"),
        ("ENAP", "Energía", "CL", "BBB"),
        ("BANCO CREDITO", "Bancos", "PE", "BBB"),
        ("CREDICORP", "Holding", "PE", "BBB"),
        ("MILPO", "Minería", "PE", "BB+"),
    ]
    durations = [0.8, 1.5, 2.3, 3.1, 3.8, 4.5, 5.4, 6.7, 7.5, 9.2, 11.5]
    ratings_grade = {"AAA": "IG", "AA+": "IG", "AA": "IG", "AA-": "IG",
                     "A+": "IG", "A": "IG", "A-": "IG",
                     "BBB+": "IG", "BBB": "IG", "BBB-": "IG",
                     "BB+": "HY", "BB": "HY", "BB-": "HY", "B+": "HY"}

    rows = []
    for n in range(1, 121):
        em = random.choice(emisores_pool)
        ticker = f"{em[0][:5].replace(' ','')}{random.choice([26,27,28,29,30,31,32,33,34,35])}"
        sector, pais, rating = em[1], em[2], em[3]
        if random.random() < 0.18:
            downgrades = {"BBB-": "BB+", "BB+": "BB", "BB": "BB-"}
            rating = downgrades.get(rating, rating)
        dur = random.choice(durations) + round(random.uniform(-0.3, 0.3), 2)
        if ratings_grade.get(rating) == "IG":
            ytm = round(random.uniform(7.2, 11.8), 2)
        else:
            ytm = round(random.uniform(10.5, 16.4), 2)
        nominal = random.choice([500, 800, 1000, 1500, 2000, 2500, 3000, 4000, 5000])
        precio = round(random.uniform(86, 104), 2)
        mv = round(nominal * precio / 100, 2)
        rows.append((ticker, em[0], sector, pais, rating, dur, ytm, nominal, precio, mv))

    for i, r in enumerate(rows, start=4):
        for j, v in enumerate(r, start=1):
            ws.cell(row=i, column=j, value=v).font = Font(color="E0E0F0", size=10)
    autosize(ws)

    ws2 = wb.create_sheet("Reglas")
    ws2["A1"] = "Notas para el ejercicio Quick Win 02"
    ws2["A1"].font = Font(bold=True, color="E85A1F", size=12)
    notas = [
        "Investment Grade = AAA hasta BBB-. High Yield = BB+ y debajo.",
        "valor_mercado_mm_cop = nominal_mm_cop × precio_clean / 100.",
        "Duración modificada se aproxima con duration_años / (1 + ytm/100).",
        "El ejercicio pide AUM total, AUM por sector, top 10 posiciones, distribución IG/HY y duración ponderada.",
        "Outliers > 2σ: revisar columnas ytm_% y precio_clean.",
    ]
    for i, n in enumerate(notas, start=3):
        ws2.cell(row=i, column=1, value=f"• {n}").font = Font(color="C0C0E0", size=10)
    ws2.column_dimensions["A"].width = 120

    wb.save(os.path.join(OUT, "02-portafolio-abril.xlsx"))


# ──────────────────────────────────────────────────────────────────────────
# 3. Prospecto emisor — Hidroeléctrica del Caribe (Word ~estructurado)
# ──────────────────────────────────────────────────────────────────────────
def make_prospecto_emisor():
    doc = Document()
    add_heading(doc, "Prospecto de Información — Bonos Hidroeléctrica del Caribe 2026", level=0, color="E85A1F")
    banner(doc, "EMISOR SINTÉTICO — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    add_heading(doc, "Resumen de la oferta", level=1)
    add_para(doc,
        "Hidroeléctrica del Caribe S.A. ESP (en adelante 'el Emisor') ofrece bonos ordinarios "
        "denominados en pesos colombianos por un monto total de COP 800.000 millones, distribuidos "
        "en tres series. La emisión ha sido inscrita en el Registro Nacional de Valores y Emisores "
        "(RNVE) y aprobada por la Superintendencia Financiera de Colombia bajo resolución "
        "0451 de marzo de 2026.")

    add_heading(doc, "1. Condiciones financieras", level=1)
    for line in [
        "Monto total: COP 800.000.000.000.",
        "Serie A (plazo 5 años): COP 250.000 M — Tasa cupón IBR + 285 bps, pago trimestral.",
        "Serie B (plazo 7 años): COP 300.000 M — Tasa cupón IPC + 415 bps, pago semestral.",
        "Serie C (plazo 10 años): COP 250.000 M — Tasa cupón fija 11,40% E.A., pago semestral.",
        "Calificación de la emisión: AA+ otorgada por Fitch Ratings Colombia el 04-mar-2026.",
        "Bookrunners: Corredores Davivienda, BTG Pactual S.A. Comisionista de Bolsa, Credicorp Capital.",
    ]:
        p = doc.add_paragraph(line, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10)

    add_heading(doc, "2. Uso de los fondos", level=1)
    add_para(doc,
        "El 62% de los recursos será destinado a refinanciar deuda bancaria sindicada con "
        "vencimiento en septiembre de 2027, el 28% financiará la expansión de la planta "
        "hidroeléctrica El Quebradón (capacidad adicional 220 MW) y el 10% restante será "
        "destinado a capital de trabajo recurrente.")

    add_heading(doc, "3. Información financiera del Emisor (cifras en COP miles de millones)", level=1)
    add_para(doc,
        "Ingresos operacionales 2023: 1.812 — 2024: 1.987 — 2025: 2.221 (CAGR 10,7%). "
        "EBITDA 2023: 712 — 2024: 824 — 2025: 949. Margen EBITDA 2025: 42,7%. "
        "Deuda financiera neta 2025: COP 2.846 mil M. Deuda neta / EBITDA: 3,0x. "
        "Cobertura de intereses (EBITDA / Gasto financiero): 4,1x. "
        "Caja y equivalentes al cierre 2025: COP 248 mil M.")

    add_heading(doc, "4. Covenants financieros (sección 4.2 del prospecto)", level=1)
    covenants = [
        "Deuda Neta / EBITDA ≤ 4,25x medido trimestralmente sobre últimos 12 meses.",
        "Cobertura de Servicio de Deuda ≥ 1,30x al cierre de cada año fiscal.",
        "Patrimonio mínimo de COP 1.000.000 millones durante toda la vida del bono.",
        "Restricción de dividendos si Deuda Neta/EBITDA > 3,75x.",
        "Cross-default con cualquier obligación superior a COP 50.000 millones.",
        "Negative pledge sobre activos productivos clave (plantas Quebradón I y II).",
        "Obligación de reporte trimestral auditado dentro de 60 días calendario.",
    ]
    for c in covenants:
        p = doc.add_paragraph(c, style="List Number")
        for r in p.runs:
            r.font.size = Pt(10)

    add_heading(doc, "5. Factores de riesgo (extracto sección 5)", level=1)
    riesgos = [
        ("Riesgo hidrológico", "Caudales históricamente variables por El Niño/La Niña. Stress test "
         "interno: caída 18% en generación durante El Niño severo."),
        ("Riesgo regulatorio", "Posible recalibración del WACC regulado por CREG entre 2026-2028."),
        ("Riesgo cambiario", "67% de la deuda financiera está en moneda local; 33% es deuda multilateral "
         "en USD con hedge parcial vía swaps."),
        ("Riesgo de contraparte", "Cobranza concentrada en 6 comercializadoras mayoristas (78% ventas)."),
        ("Riesgo de construcción", "Ampliación El Quebradón estimada en USD 145 M con CAPEX overrun "
         "contingencia 12%."),
    ]
    for tit, txt in riesgos:
        add_para(doc, tit, bold=True, color="E85A1F")
        add_para(doc, txt)

    add_heading(doc, "6. Comparables sectoriales", level=1)
    add_para(doc,
        "ISA 2032 cupón 8,90% spread vs TES 200 bps. Promigas 2030 cupón 9,75% spread 245 bps. "
        "Celsia 2028 cupón IPC + 380 bps. Esta emisión Serie C 11,40% E.A. equivale a spread "
        "aproximado de 295 bps sobre TES 10 años (referencia abril 2026).")

    add_heading(doc, "7. Calendario de la emisión", level=1)
    for t in [
        "Apertura del libro: 22-abr-2026, 9:00 am.",
        "Cierre del libro: 23-abr-2026, 12:00 m.",
        "Asignación: 23-abr-2026, 3:00 pm.",
        "Cumplimiento: 28-abr-2026 (T+3).",
        "Primer pago de intereses: 28-jul-2026.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")

    doc.save(os.path.join(OUT, "03-prospecto-hidrocaribe.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 4. KIID paragraph EN — para traducir con DeepSeek vs ChatGPT
# ──────────────────────────────────────────────────────────────────────────
def make_kiid_en():
    doc = Document()
    add_heading(doc, "Key Investor Information Document — Sample Paragraph (EN)", level=0)
    banner(doc, "SAMPLE / SYNTHETIC — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    add_heading(doc, "Instructions", level=1)
    add_para(doc,
        "Use the paragraph below for the translation comparison task between DeepSeek, "
        "ChatGPT and Kimi. The goal is to evaluate which engine preserves financial precision "
        "and Colombian Spanish technicality.",
        italic=True, size=10)

    add_heading(doc, "Paragraph 1 — Investment objective and policy", level=1)
    add_para(doc,
        "The Sub-Fund seeks to deliver long-term capital appreciation primarily through exposure "
        "to senior secured fixed-rate corporate bonds issued by Latin American utilities, with a "
        "modified duration target between 4.0 and 6.5 years. Up to 25% of net asset value may be "
        "allocated to subordinated or hybrid debt instruments, including AT1 perpetual notes, "
        "provided the issuer maintains an investment-grade local-scale rating. Derivative instruments "
        "may be used for hedging foreign exchange exposure and managing duration; their gross notional "
        "shall not exceed 100% of NAV. The Sub-Fund is actively managed and references the iBoxx LatAm "
        "Utilities USD Investment Grade Index for performance comparison purposes only.")

    add_heading(doc, "Paragraph 2 — Risk warning", level=1)
    add_para(doc,
        "Investors should be aware that the Sub-Fund is subject to credit, interest rate, currency, "
        "liquidity and concentration risks. The use of leverage, where permitted up to 110% of NAV "
        "via repurchase agreements, may magnify losses during periods of market stress. Bid-ask spreads "
        "on LatAm corporate debt can widen materially in stressed conditions, potentially impairing the "
        "Manager's ability to realise assets at quoted prices.")

    add_heading(doc, "Paragraph 3 — Charges", level=1)
    add_para(doc,
        "The Sub-Fund applies a 1.10% per annum management fee on net assets, accrued daily, plus a "
        "performance fee of 15% over the high-water mark of the reference index, subject to a hurdle "
        "rate of 250 basis points above the 1-month USD SOFR.")

    add_heading(doc, "Suggested prompt", level=1)
    add_para(doc,
        "“Traduce este KIID al español manteniendo precisión financiera, citas regulatorias colombianas "
        "(SFC) donde aplique, y unidad de medida en pesos colombianos cuando sea relevante para un "
        "inversionista institucional local.”",
        italic=True)

    doc.save(os.path.join(OUT, "04-kiid-en-paragraphs.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 5. Serie de retornos para VaR — Ollama / DeepSeek local
# ──────────────────────────────────────────────────────────────────────────
def make_serie_var():
    wb = Workbook()
    ws = wb.active
    ws.title = "Retornos"

    ws["A1"] = "Serie sintética de retornos diarios — Portafolio BTG LatAm (2 años hábiles)"
    ws["A1"].font = Font(bold=True, color="E85A1F", size=12)
    ws.merge_cells("A1:C1")
    ws["A2"] = "Útil para el ejercicio Quick Win 05 (DeepSeek-R1 sobre Ollama — VaR histórico)."
    ws["A2"].font = Font(italic=True, color="C0C0E0", size=10)

    headers = ["fecha", "retorno_diario_%", "nav_index"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(headers))

    fecha = datetime(2024, 4, 22)
    nav = 100.0
    row = 5
    for _ in range(504):
        if fecha.weekday() >= 5:
            fecha += timedelta(days=1)
            continue
        shock = random.gauss(0, 1)
        if random.random() < 0.04:
            shock *= 3.2
        r = round(0.025 + shock * 0.42, 4)
        nav *= (1 + r / 100)
        ws.cell(row=row, column=1, value=fecha.strftime("%Y-%m-%d"))
        ws.cell(row=row, column=2, value=r)
        ws.cell(row=row, column=3, value=round(nav, 4))
        fecha += timedelta(days=1)
        row += 1
    autosize(ws)

    ws2 = wb.create_sheet("Prompt sugerido")
    ws2["A1"] = (
        "Prompt para deepseek-r1:7b en Ollama"
    )
    ws2["A1"].font = Font(bold=True, color="E85A1F", size=12)
    bloque = [
        "Carga la serie de retornos del CSV/Excel adjunto (columna retorno_diario_%).",
        "Escribe una función Python `var_historico(retornos: list[float], conf: float = 0.95)` que:",
        "  - levante ValueError si len(retornos) < 30,",
        "  - calcule el percentil (1 - conf) sobre la serie ordenada,",
        "  - devuelva el VaR con signo negativo (pérdida).",
        "Incluye docstring estilo Google, tipo hints y un bloque if __name__ == '__main__' que "
        "calcule VaR al 95% y al 99% con la serie completa.",
        "No uses pandas — solo statistics o numpy si está disponible.",
    ]
    for i, t in enumerate(bloque, start=3):
        ws2.cell(row=i, column=1, value=t).font = Font(color="C0C0E0", size=10)
    ws2.column_dimensions["A"].width = 110

    wb.save(os.path.join(OUT, "05-serie-retornos-var.xlsx"))


# ──────────────────────────────────────────────────────────────────────────
# 6. Circular SFC sintética — para TL;DR
# ──────────────────────────────────────────────────────────────────────────
def make_circular_sfc():
    doc = Document()
    add_heading(doc, "Circular Externa 008 de 2026 — Superintendencia Financiera de Colombia", level=0, color="E85A1F")
    banner(doc, "CIRCULAR SINTÉTICA — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    add_heading(doc, "Asunto", level=1)
    add_para(doc,
        "Modificación a la Parte II, Título II, Capítulo XX del Estatuto Orgánico del Sistema "
        "Financiero, en lo referente al cómputo de la relación de solvencia adicional para "
        "establecimientos de crédito que operen con instrumentos derivados sobre tasa de interés.")

    add_heading(doc, "Fecha de entrada en vigor", level=1)
    add_para(doc, "1° de julio de 2026, con periodo de transición de 90 días calendario.")

    add_heading(doc, "Artículo 1 — Objeto", level=1)
    add_para(doc,
        "La presente Circular establece el ajuste al factor de ponderación de riesgo aplicable a las "
        "exposiciones netas en derivados sobre tasa de interés celebrados con contraparte central "
        "(CCP) y no contraparte central, dentro del cómputo de los activos ponderados por nivel de "
        "riesgo de crédito (APNR). El factor pasa de 4% a 2% para exposiciones con CCP y de 8% a "
        "10% para exposiciones bilaterales no compensadas, alineándose con el marco de Basilea III "
        "revisado en su versión final 2023.")

    add_heading(doc, "Artículo 2 — Reporte de información", level=1)
    add_para(doc,
        "Las entidades sometidas a inspección y vigilancia deberán remitir a esta Superintendencia, "
        "a más tardar el décimo quinto día hábil de cada mes, el formato 524 actualizado con la "
        "información de exposiciones netas por contraparte, segregando CCP y no-CCP, en pesos "
        "colombianos y moneda extranjera, de acuerdo con la matriz anexa a la presente Circular.")

    add_heading(doc, "Artículo 3 — Régimen de transición", level=1)
    add_para(doc,
        "Las exposiciones existentes al 30 de junio de 2026 podrán continuar bajo el régimen "
        "anterior hasta su vencimiento o el 31 de diciembre de 2026, lo que ocurra primero. Toda "
        "operación nueva celebrada desde el 1° de julio de 2026 deberá aplicar inmediatamente los "
        "nuevos factores.")

    add_heading(doc, "Artículo 4 — Sanciones", level=1)
    add_para(doc,
        "El incumplimiento de las obligaciones de reporte y cómputo será sancionable conforme al "
        "Estatuto Orgánico del Sistema Financiero, sin perjuicio de las acciones correctivas que "
        "esta Superintendencia disponga.")

    add_heading(doc, "Artículo 5 — Disposiciones finales", level=1)
    add_para(doc,
        "Esta Circular deroga las disposiciones que le sean contrarias, en especial los numerales "
        "3.2.1 y 3.2.2 de la Circular Externa 014 de 2023.")

    add_heading(doc, "Anexo técnico — Ejemplo de cálculo", level=1)
    add_para(doc,
        "Banco hipotético con exposición neta de COP 1.200.000 M en swaps de tasa contra CCP: "
        "APNR previo = 1.200.000 × 0,04 = COP 48.000 M. APNR nuevo = 1.200.000 × 0,02 = COP 24.000 M. "
        "Impacto en solvencia: liberación de capital regulatorio de aproximadamente COP 2.640 M "
        "(asumiendo solvencia objetivo del 11%).")

    add_heading(doc, "Prompt sugerido para DeepSeek", level=1)
    add_para(doc,
        "“Eres el chief of staff de un MD en banca. Resume esta circular SFC en 5 bullets de máximo "
        "22 palabras cada uno, para una reunión con la junta en 1 hora. Incluye fecha de entrada en "
        "vigor, impacto directo para un banco universal colombiano y 1 acción recomendada.”",
        italic=True)

    doc.save(os.path.join(OUT, "06-circular-sfc-008-2026.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 7. TES B 2032 — datos para benchmark triple (Excel)
# ──────────────────────────────────────────────────────────────────────────
def make_tes_b_2032():
    wb = Workbook()
    ws = wb.active
    ws.title = "TES B 2032"

    ws["A1"] = "TES B 2032 · Datos de referencia para benchmark triple (Quick Win 07)"
    ws["A1"].font = Font(bold=True, color="E85A1F", size=12)
    ws.merge_cells("A1:E1")

    info = [
        ("ISIN", "COTFB022032"),
        ("Emisor", "Gobierno Nacional de Colombia"),
        ("Cupón anual", "7.25%"),
        ("Base", "365 días"),
        ("Fecha emisión", "2022-07-26"),
        ("Vencimiento", "2032-07-26"),
        ("Precio sucio", "94.85"),
        ("Precio limpio", "92.50"),
        ("Plazo residual aprox", "6.3 años"),
        ("Frecuencia cupón", "Anual"),
        ("Calificación", "BBB (S&P) / Baa2 (Moody's)"),
    ]
    for i, (k, v) in enumerate(info, start=3):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True, color="00D4E5", size=11)
        ws.cell(row=i, column=2, value=v).font = Font(color="E0E0F0", size=11)

    # Flujo de caja
    ws.cell(row=15, column=1, value="Flujo de caja proyectado (nominal 100)")
    ws.cell(row=15, column=1).font = Font(bold=True, color="E85A1F", size=11)
    ws.cell(row=16, column=1, value="periodo_años")
    ws.cell(row=16, column=2, value="flujo_cupon")
    ws.cell(row=16, column=3, value="flujo_principal")
    ws.cell(row=16, column=4, value="flujo_total")
    style_header(ws, 16, 4)

    for i in range(1, 7):
        ws.cell(row=16 + i, column=1, value=i)
        ws.cell(row=16 + i, column=2, value=7.25)
        ws.cell(row=16 + i, column=3, value=100 if i == 6 else 0)
        ws.cell(row=16 + i, column=4, value=7.25 + (100 if i == 6 else 0))

    ws2 = wb.create_sheet("Glosario")
    items = [
        ("Duración modificada", "Sensibilidad de precio ante cambio de 1% en tasa: -ΔP/P ≈ -D_mod × Δy."),
        ("Convexidad", "Segunda derivada del precio respecto a la tasa: ajuste de la aproximación lineal de duración."),
        ("YTM aproximada", "TIR del flujo, asumiendo reinversión a esa misma tasa hasta vencimiento."),
        ("Base 365", "Cómputo de tiempos sobre 365 días; convención local en TES."),
    ]
    for i, (k, v) in enumerate(items, start=2):
        ws2.cell(row=i, column=1, value=k).font = Font(bold=True, color="00D4E5", size=11)
        ws2.cell(row=i, column=2, value=v).font = Font(color="E0E0F0", size=10)
    ws2.column_dimensions["A"].width = 26
    ws2.column_dimensions["B"].width = 100
    autosize(ws)

    wb.save(os.path.join(OUT, "07-tes-b-2032-benchmark.xlsx"))


# ──────────────────────────────────────────────────────────────────────────
# 8. Memoria anual sintética — Ecocemento Andino 2025 (Kimi long context)
# ──────────────────────────────────────────────────────────────────────────
def make_memoria_anual():
    doc = Document()
    add_heading(doc, "Memoria Anual 2025 — Ecocemento Andino S.A.", level=0, color="E85A1F")
    banner(doc, "EMISOR SINTÉTICO — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    # Carta del CEO
    add_heading(doc, "1. Carta del CEO", level=1)
    add_para(doc,
        "Estimados accionistas, 2025 fue un año de consolidación operativa. Cerramos con ventas "
        "de 7,1 millones de toneladas de cemento (+4,8% a/a), un EBITDA de COP 1.842 mil millones "
        "(+11,2% a/a) y una utilidad neta de COP 612 mil millones. Nuestra estrategia 'Andino "
        "Verde 2030' avanzó con tres hitos: clinker descarbonizado en planta Sogamoso, alianza "
        "con Holcim para captura de CO₂ y emisión de bono verde local por COP 600.000 millones "
        "con sobre-demanda de 2,8x.")

    add_heading(doc, "2. Desempeño financiero", level=1)
    add_para(doc,
        "Ingresos consolidados: COP 4.872 mil M (+9,1%). Costo de ventas: COP 3.130 mil M. "
        "Margen bruto: 35,8% (vs 33,4% en 2024). EBITDA: COP 1.842 mil M, margen 37,8%. "
        "Utilidad operacional: COP 1.121 mil M. Resultado financiero neto: −COP 312 mil M. "
        "Utilidad neta: COP 612 mil M. UPA: COP 1.842.")
    add_para(doc,
        "Inversión de capital (Capex): COP 745 mil M (15,3% de ingresos), concentrada en (i) "
        "expansión planta Sogamoso COP 412 mil M, (ii) modernización molinos planta Yumbo COP 198 "
        "mil M, (iii) infraestructura logística pacífico COP 135 mil M.")

    add_heading(doc, "3. Operaciones por país", level=1)
    paises = [
        ("Colombia", "5,2 Mt despachadas. Margen EBITDA 38,9%. Capex foco Sogamoso descarbonizado."),
        ("Perú", "1,1 Mt. Margen EBITDA 32,4%. Negociación colectiva cerrada sin paro."),
        ("Ecuador", "0,5 Mt. Margen EBITDA 28,1%. Recuperación tras crisis fiscal 2024."),
        ("Panamá", "0,3 Mt. Margen EBITDA 41,8%. Proyectos canal + infraestructura urbana."),
    ]
    for k, v in paises:
        add_para(doc, k, bold=True, color="E85A1F")
        add_para(doc, v)

    add_heading(doc, "4. Sostenibilidad — Andino Verde 2030", level=1)
    add_para(doc,
        "Reducción emisiones específicas: 612 kg CO₂/t cemento (−6,3% a/a). Meta 2030: 480 kg/t. "
        "Energía renovable en operaciones: 47% (PPA solar Yumbo + autogeneración). "
        "Agua reciclada en procesos: 78%. Inversión social comunidades: COP 28 mil M.")

    add_heading(doc, "5. Gobierno corporativo", level=1)
    add_para(doc,
        "Junta directiva ampliada de 7 a 9 miembros, con cuatro independientes (44%). Nueva política "
        "de remuneración ejecutiva con 30% del bono ligado a indicadores ESG verificables. "
        "Compliance: cero incidentes materiales reportados a SFC durante el año.")

    add_heading(doc, "6. Factores de riesgo 2025", level=1)
    riesgos_2025 = [
        ("Operacional", "Disponibilidad de clinker tras paro 12 días en planta Cali (mitigado con stocks)."),
        ("Operacional", "Costo logístico cabotaje pacífico +18% por tarifas portuarias."),
        ("Financiero", "Sensibilidad a tasa flotante: 42% deuda IBR; cada 100 bps = COP 28 mil M Ebit."),
        ("Regulatorio", "Implementación impuesto al carbono Fase II: costo estimado COP 22 mil M."),
        ("Estratégico", "Competencia importadora de cemento desde Asia, presión sobre precios premium."),
        ("Cambio climático", "Eventos hídricos extremos afectaron extracción de áridos Q2."),
        ("Reputacional", "Demanda colectiva pendiente por emisión histórica polvo planta Yumbo (provisión COP 12 mil M)."),
        ("Ciberseguridad", "Intento de ransomware en filial Perú; sin pérdida de datos, contenido en 6 horas."),
        ("Talento", "Rotación senior 11,2% (vs 8,4% en 2024)."),
        ("Materialización 2025", "Riesgo operacional Cali se materializó parcialmente; pérdida de margen estimada COP 18 mil M (mitigada)."),
    ]
    for cat, txt in riesgos_2025:
        add_para(doc, f"[{cat}] {txt}")

    add_heading(doc, "7. Comparación de riesgos 2024 vs 2025", level=1)
    for t in [
        "Nuevos en 2025: impuesto al carbono Fase II, demanda colectiva polvo Yumbo, intento ransomware.",
        "Salieron del top 10 en 2025: riesgo FX sobre deuda en USD (cobertura completada en agosto).",
        "Cambio de orden: riesgo logístico subió del puesto 7 al 2 por tarifas portuarias.",
        "Mantienen materialidad: hidrológico, regulatorio, talento.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10)

    add_heading(doc, "8. Estrategia 2026-2030", level=1)
    add_para(doc,
        "Cinco palancas: (1) descarbonización clinker — meta 480 kg CO₂/t a 2030; (2) expansión "
        "Centroamérica orgánica + M&A selectiva; (3) cemento especializado para infraestructura "
        "vial 4G; (4) logística marítima propia ruta pacífico; (5) optimización capital de trabajo "
        "para liberar COP 180 mil M en tres años.")

    add_heading(doc, "9. Estados financieros — resumen", level=1)
    add_para(doc,
        "Activo total: COP 8.412 mil M. Pasivo total: COP 5.230 mil M. Patrimonio: COP 3.182 mil M. "
        "Deuda financiera bruta: COP 3.110 mil M. Deuda financiera neta: COP 2.642 mil M. "
        "Deuda neta / EBITDA: 1,43x. Cobertura intereses: 5,9x.")

    add_heading(doc, "10. Hechos relevantes posteriores", level=1)
    add_para(doc,
        "El 14 de febrero de 2026 se aprobó la emisión del bono verde local. El 02 de marzo de "
        "2026 se firmó el acuerdo de captura de CO₂ con Holcim. El 19 de marzo de 2026, S&P "
        "ratificó la calificación BBB con perspectiva estable.")

    add_heading(doc, "Prompts sugeridos (Quick Win 08)", level=1)
    for t in [
        "1) Extrae los 10 factores de riesgo declarados en sección 6 con página, categoría y materialización.",
        "2) Diff de riesgos 2024 vs 2025 según sección 7 — qué desapareció y qué es nuevo.",
        "3) Tres preguntas que haría un analista senior al CFO basándose en estos cambios.",
    ]:
        add_para(doc, t, italic=True)

    doc.save(os.path.join(OUT, "08-memoria-ecocemento-2025.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 9. Balance trimestral — Ratios financieros (Word con tablas)
# ──────────────────────────────────────────────────────────────────────────
def make_balance_trimestral():
    doc = Document()
    add_heading(doc, "Balance Trimestral Q1 2026 — TextilAndes S.A.", level=0, color="E85A1F")
    banner(doc, "EMISOR SINTÉTICO — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    add_para(doc,
        "Balance de un emisor corporativo colombiano de tamaño mediano. Cifras en COP miles de "
        "millones. Datos sintéticos para el ejercicio Quick Win 09 (DeepSeek file upload + DeepThink).",
        italic=True, size=10)

    add_heading(doc, "Estado de Situación Financiera", level=1)
    tabla1 = [
        ("Línea", "Mar-26", "Dic-25"),
        ("ACTIVO CORRIENTE", "", ""),
        ("Efectivo y equivalentes", "182", "215"),
        ("Cuentas por cobrar comerciales", "418", "385"),
        ("Inventarios", "612", "580"),
        ("Otros activos corrientes", "94", "88"),
        ("Total activo corriente", "1.306", "1.268"),
        ("ACTIVO NO CORRIENTE", "", ""),
        ("Propiedad, planta y equipo, neto", "2.142", "2.110"),
        ("Intangibles + Goodwill", "486", "486"),
        ("Otros activos no corrientes", "118", "112"),
        ("Total activo no corriente", "2.746", "2.708"),
        ("ACTIVO TOTAL", "4.052", "3.976"),
        ("PASIVO CORRIENTE", "", ""),
        ("Cuentas por pagar comerciales", "298", "276"),
        ("Deuda financiera corriente", "412", "390"),
        ("Otros pasivos corrientes", "186", "174"),
        ("Total pasivo corriente", "896", "840"),
        ("PASIVO NO CORRIENTE", "", ""),
        ("Deuda financiera no corriente", "1.214", "1.245"),
        ("Pasivos por impuesto diferido", "168", "160"),
        ("Otros pasivos no corrientes", "92", "88"),
        ("Total pasivo no corriente", "1.474", "1.493"),
        ("PASIVO TOTAL", "2.370", "2.333"),
        ("PATRIMONIO TOTAL", "1.682", "1.643"),
    ]
    t = doc.add_table(rows=len(tabla1), cols=3)
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(tabla1):
        for j, v in enumerate(row):
            c = t.cell(i, j)
            c.text = v
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10)
                if i == 0 or row[0] in ("ACTIVO CORRIENTE", "ACTIVO NO CORRIENTE",
                                         "PASIVO CORRIENTE", "PASIVO NO CORRIENTE",
                                         "ACTIVO TOTAL", "PASIVO TOTAL", "PATRIMONIO TOTAL"):
                    r.bold = True

    add_heading(doc, "Estado de Resultados — Q1 2026 (3 meses)", level=1)
    tabla2 = [
        ("Línea", "Q1 2026", "Q1 2025"),
        ("Ingresos operacionales", "742", "692"),
        ("Costo de ventas", "(498)", "(471)"),
        ("Utilidad bruta", "244", "221"),
        ("Gastos administración y ventas", "(118)", "(112)"),
        ("Depreciación y amortización", "(56)", "(54)"),
        ("EBIT", "70", "55"),
        ("Gasto financiero neto", "(28)", "(27)"),
        ("Utilidad antes de impuestos", "42", "28"),
        ("Impuesto a la renta", "(13)", "(9)"),
        ("Utilidad neta", "29", "19"),
        ("EBITDA", "126", "109"),
    ]
    t2 = doc.add_table(rows=len(tabla2), cols=3)
    t2.style = "Light Grid Accent 1"
    for i, row in enumerate(tabla2):
        for j, v in enumerate(row):
            c = t2.cell(i, j)
            c.text = v
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10)
                if i == 0 or row[0] in ("EBIT", "EBITDA", "Utilidad neta", "Utilidad bruta"):
                    r.bold = True

    add_heading(doc, "Notas relevantes", level=1)
    notas = [
        "Línea 'Deuda financiera no corriente' incluye bono local 9,25% vencimiento 2029 por COP 800 mil M.",
        "Cobertura SWAP sobre 60% de la deuda corriente.",
        "Provisión riesgo crediticio sobre cartera: 3,2% del saldo bruto.",
        "Inventarios valorados a costo promedio ponderado.",
    ]
    for n in notas:
        p = doc.add_paragraph(n, style="List Bullet")

    add_heading(doc, "Prompt sugerido (Quick Win 09)", level=1)
    add_para(doc,
        "“De este balance calcula: razón corriente, prueba ácida, endeudamiento total, deuda "
        "financiera neta / EBITDA, cobertura de intereses, ROE, ROA, margen EBITDA. Devuelve tabla "
        "markdown con valor, fórmula y la línea exacta del balance de donde viene cada componente.”",
        italic=True)

    doc.save(os.path.join(OUT, "09-balance-textilandes-q1.docx"))


# ──────────────────────────────────────────────────────────────────────────
# 10. Brief de DD express — emisor desconocido (Word + plantilla)
# ──────────────────────────────────────────────────────────────────────────
def make_dd_brief():
    doc = Document()
    add_heading(doc, "Brief DD Express — Emisor candidato 'AgroValle Pacífico S.A.'", level=0, color="E85A1F")
    banner(doc, "EMISOR SINTÉTICO — Insumo didáctico Sesión 6 · BTG × NODO EAFIT")

    add_heading(doc, "Contexto del ejercicio", level=1)
    add_para(doc,
        "Vas a tu primera llamada con AgroValle Pacífico S.A. mañana 9 am. Es un emisor "
        "agroindustrial colombiano (caña, etanol, energía cogeneración) que hasta hoy no conocías. "
        "Tu jefe te pidió un dossier ejecutivo y cinco preguntas punzantes. Usa este brief como "
        "punto de partida y completa con web search en DeepSeek y Kimi.")

    add_heading(doc, "Datos básicos (públicos)", level=1)
    datos = [
        ("Sector", "Agroindustria — caña de azúcar, etanol, cogeneración eléctrica"),
        ("Ubicación principal", "Valle del Cauca, Colombia"),
        ("Años en operación", "47"),
        ("Empleados", "~3.200"),
        ("Ingresos 2024 (estimado)", "COP 1.480 mil millones"),
        ("EBITDA 2024 (estimado)", "COP 312 mil millones"),
        ("Calificación local", "AA– (Fitch Ratings Colombia, revisión sep-2025)"),
        ("Emisiones vigentes", "Bono local 2027 cupón 8,75% E.A. · monto vigente COP 280 mil M"),
        ("Eventos materiales 12m", "Cambio de CFO (ene-26) · Inversión nuevo trapiche (40 M USD anunciado mar-26)"),
        ("Mediático", "Demanda ambiental por captación río pendiente desde 2024"),
        ("Web search sugerida", "site:bvc.com.co AgroValle · site:fitchratings.com 'AgroValle Pacífico'"),
    ]
    for k, v in datos:
        add_para(doc, k, bold=True, color="00D4E5")
        add_para(doc, v)

    add_heading(doc, "Plantilla del dossier (a llenar)", level=1)
    add_para(doc,
        "1) 5 hechos materiales últimos 12 meses (cada uno con URL verificable).\n"
        "2) Rating actual + última revisión.\n"
        "3) Cambios en gobierno corporativo.\n"
        "4) 3 riesgos rojos identificados.\n"
        "5) 5 preguntas punzantes para el CFO, ordenadas por materialidad.",
        size=11)

    add_heading(doc, "Prompt sugerido (Quick Win 10)", level=1)
    add_para(doc,
        "“Dossier ejecutivo del emisor AgroValle Pacífico S.A. a abril 2026: eventos materiales "
        "últimos 12 meses, rating vigente (S&P/Moody/Fitch), últimas revisiones, cambios en "
        "gobierno corporativo y noticias negativas si existen. Cita URL exacta de cada dato.”",
        italic=True)

    add_heading(doc, "Checklist de verificación", level=1)
    for t in [
        "Click en cada URL devuelta y comprueba que existe y mantiene la cifra.",
        "Cruza un mismo hecho material en DeepSeek y Kimi; descarta si solo aparece en uno.",
        "Confirma rating en página oficial del calificador, no en blog secundario.",
        "Si un evento no aparece en BVC ni en SFC, marca con (⚠ no oficial) en tu dossier.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")

    doc.save(os.path.join(OUT, "10-dd-brief-agrovalle.docx"))


# ──────────────────────────────────────────────────────────────────────────
# Run all
# ──────────────────────────────────────────────────────────────────────────
def main():
    make_utilities_brief()
    make_portafolio_abril()
    make_prospecto_emisor()
    make_kiid_en()
    make_serie_var()
    make_circular_sfc()
    make_tes_b_2032()
    make_memoria_anual()
    make_balance_trimestral()
    make_dd_brief()
    print(f"OK — archivos generados en {OUT}")
    for f in sorted(os.listdir(OUT)):
        p = os.path.join(OUT, f)
        print(f"  · {f} ({os.path.getsize(p)/1024:.1f} KB)")


if __name__ == "__main__":
    main()
