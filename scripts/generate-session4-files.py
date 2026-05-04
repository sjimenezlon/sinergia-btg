"""
Genera los archivos descargables de la Sesión 4 — datos sintéticos para que los equipos
los carguen a Claude / ChatGPT / Gemini / NotebookLM / DeepSeek / Mistral / Kimi.
Ningún dato real de BTG ni de clientes — todo es sintético y verosímil.
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "..", "public", "sesion-4")
os.makedirs(OUT, exist_ok=True)


# ──────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────
HEADER_FILL = PatternFill("solid", fgColor="0D1229")
HEADER_FONT = Font(bold=True, color="E85A1F", size=11)
SUB_FILL = PatternFill("solid", fgColor="151A3A")
THIN = Side(border_style="thin", color="2A2F4F")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_header(ws, row, cols):
    for col in range(1, cols + 1):
        c = ws.cell(row=row, column=col)
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = BORDER


def autosize(ws):
    from openpyxl.cell.cell import MergedCell
    widths = {}
    for row in ws.iter_rows():
        for c in row:
            if isinstance(c, MergedCell) or c.value is None:
                continue
            widths[c.column_letter] = max(widths.get(c.column_letter, 10), len(str(c.value)) + 2)
    for letter, w in widths.items():
        ws.column_dimensions[letter].width = min(w, 60)


def add_heading(doc, text, level=1, color="E85A1F"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


# ──────────────────────────────────────────────────────────────────────────
# 1. Cementos Portales — CIM (Word)
# ──────────────────────────────────────────────────────────────────────────
def make_cim_word():
    doc = Document()
    add_heading(doc, "Cementos Portales S.A.", level=0, color="E85A1F")
    p = doc.add_paragraph()
    p.add_run("Confidential Information Memorandum · Q1 2026 · Datos sintéticos para ejercicio educativo").italic = True

    add_heading(doc, "1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Cementos Portales (en adelante, 'la Compañía') es un productor regional de cemento gris y mortero "
        "con operaciones en Colombia, Perú y Ecuador. Fundada en 1978 en Medellín, opera 3 plantas integradas "
        "y 14 centros de distribución. Capacidad instalada: 2.4 Mt/año. Utilización 2025: 71%."
    )
    doc.add_paragraph(
        "El accionista controlante (Familia Portales, 62%) ha mandatado la búsqueda de un socio estratégico "
        "o comprador del 100% para apoyar la expansión a Centroamérica y refinanciar deuda corporativa."
    )

    add_heading(doc, "2. Performance financiero (cifras en COP miles de millones)", level=1)
    table = doc.add_table(rows=8, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(["Concepto", "2022", "2023", "2024", "2025"]):
        hdr[i].text = t
    rows = [
        ("Ingresos", "612", "684", "729", "751"),
        ("EBITDA", "98", "108", "112", "104"),
        ("Margen EBITDA", "16.0%", "15.8%", "15.4%", "13.8%"),
        ("Deuda neta", "298", "312", "336", "356"),
        ("Net Debt / EBITDA", "3.04x", "2.89x", "3.00x", "3.42x"),
        ("DSO (días)", "45", "52", "65", "78"),
        ("Capex", "32", "41", "38", "29"),
    ]
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            table.rows[i].cells[j].text = v

    add_heading(doc, "3. Riesgos identificados por el management", level=1)
    riesgos = [
        ("Mercado", "Caída de demanda residencial en Lima (-12% YoY) por desaceleración del sector construcción."),
        ("Operacional", "Una de las plantas (Cundinamarca) opera con horno de 1996 — eficiencia energética 18% por debajo de pares regionales."),
        ("Financiero", "Covenant Net Debt/EBITDA con Bancolombia: 3.50x. Cierre 2025: 3.42x. Margen: 0.08x."),
        ("Regulatorio", "DIAN notificó ajuste fiscal por COP 8.200 millones el 12-feb-2026. La Compañía no provisionó."),
        ("Reputacional", "Demanda colectiva pendiente en Loja (Ecuador) por contaminación de aguas. Provisión: COP 1.500M."),
    ]
    for tipo, desc in riesgos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{tipo}: ")
        r.bold = True
        p.add_run(desc)

    add_heading(doc, "4. Tesis del banquero asesor", level=1)
    doc.add_paragraph(
        "Los múltiplos de transacciones recientes en cemento LatAm (mediana EV/EBITDA: 8.2x; rango: 6.1x–11.3x) "
        "implican un Enterprise Value entre COP 634mm y COP 1.176mm. La Compañía requiere ajustes por: "
        "(i) covenant tight con Bancolombia, (ii) contingencia DIAN no provisionada, "
        "(iii) gap de eficiencia en planta Cundinamarca."
    )

    add_heading(doc, "5. Información que se entrega en data room", level=1)
    docs = [
        "Estados financieros auditados 2022–2025 (KPMG)",
        "Modelo financiero Excel con proyecciones 2026–2030",
        "Term sheet propuesto (sección 4.2 menciona covenant trigger)",
        "Legal opinion de Brigard Urrutia (sección 7: contingencia DIAN)",
        "10 contratos materiales con clientes top",
        "Audits ambientales 2023–2025",
        "Reportes de operación mensuales por planta",
    ]
    for d in docs:
        doc.add_paragraph(d, style="List Number")

    add_heading(doc, "6. Cronograma transaccional", level=1)
    cronograma = [
        ("Lanzamiento de teaser", "15-may-2026"),
        ("Entrada al data room (Round 1)", "01-jun-2026"),
        ("Entrega de NBO (Non-Binding Offer)", "01-jul-2026"),
        ("Management presentations", "15-jul-2026"),
        ("Entrega de Binding Offer", "30-ago-2026"),
        ("Cierre target", "30-nov-2026"),
    ]
    table = doc.add_table(rows=len(cronograma) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Hito"
    table.rows[0].cells[1].text = "Fecha"
    for i, (h, f) in enumerate(cronograma, start=1):
        table.rows[i].cells[0].text = h
        table.rows[i].cells[1].text = f

    p = doc.add_paragraph()
    p.add_run(
        "\n— Documento sintético para ejercicio educativo. Cualquier semejanza con compañías reales es coincidencia. —"
    ).italic = True

    doc.save(os.path.join(OUT, "01-cementos-portales-cim.docx"))
    print("✓ 01-cementos-portales-cim.docx")


# ──────────────────────────────────────────────────────────────────────────
# 2. DCF Emisores Colcap (Excel)
# ──────────────────────────────────────────────────────────────────────────
def make_dcf_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Assumptions"

    ws["A1"] = "DCF Assumptions · 5 emisores Colcap · Q1 2026"
    ws["A1"].font = Font(bold=True, size=14, color="E85A1F")
    ws.merge_cells("A1:H1")

    ws["A2"] = "Datos sintéticos · usar como input para Code Interpreter"
    ws["A2"].font = Font(italic=True, color="888888")
    ws.merge_cells("A2:H2")

    headers = ["Ticker", "Sector", "Revenue 2025 (COP B)", "EBITDA 2025 (COP B)",
               "Margen EBITDA", "WACC base", "g terminal", "Net Debt / EBITDA"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(headers))

    issuers = [
        ("ECOPETROL", "Energía", 152340, 56340, 0.37, 0.115, 0.025, 1.85),
        ("BANCOLOMBIA", "Bancos", 28450, 9420, 0.33, 0.108, 0.030, 0.0),
        ("ISA", "Energía / Infra", 12780, 5340, 0.42, 0.092, 0.025, 2.10),
        ("CEMARGOS", "Materiales", 14250, 2380, 0.17, 0.118, 0.025, 2.65),
        ("GRUPOSURA", "Servicios fin.", 9620, 2480, 0.26, 0.100, 0.030, 0.0),
    ]
    for i, row in enumerate(issuers, start=5):
        for j, v in enumerate(row, start=1):
            c = ws.cell(row=i, column=j, value=v)
            c.fill = SUB_FILL
            c.font = Font(color="FFFFFF")
            c.border = BORDER

    ws.cell(row=12, column=1, value="Ranges de sensibilidad").font = Font(bold=True, color="E85A1F")
    ws.cell(row=13, column=1, value="WACC")
    ws.cell(row=13, column=2, value="−200 bps")
    ws.cell(row=13, column=3, value="−100 bps")
    ws.cell(row=13, column=4, value="base")
    ws.cell(row=13, column=5, value="+100 bps")
    ws.cell(row=13, column=6, value="+200 bps")
    ws.cell(row=14, column=1, value="g terminal")
    ws.cell(row=14, column=2, value="2.0%")
    ws.cell(row=14, column=3, value="2.5%")
    ws.cell(row=14, column=4, value="3.0%")
    ws.cell(row=14, column=5, value="3.5%")
    ws.cell(row=14, column=6, value="4.0%")

    ws.cell(row=16, column=1, value="Tarea sugerida para el agente:").font = Font(bold=True, color="00E5A0")
    tarea = [
        "1. Para cada emisor, proyectar EBITDA 5 años con CAGR sectorial.",
        "2. Calcular Free Cash Flow asumiendo Capex = 60% Depreciación.",
        "3. Construir matriz 5x5 de Enterprise Value (WACC × g).",
        "4. Identificar el emisor con mayor sensibilidad y generar tornado chart.",
        "5. Devolver tabla resumen con EV/EBITDA implícito por emisor.",
    ]
    for i, t in enumerate(tarea, start=17):
        ws.cell(row=i, column=1, value=t)

    autosize(ws)

    # Hoja 2: histórico
    ws2 = wb.create_sheet("Histórico")
    headers2 = ["Año", "ECOPETROL EBITDA", "BANCOLOMBIA EBITDA", "ISA EBITDA", "CEMARGOS EBITDA", "GRUPOSURA EBITDA"]
    for i, h in enumerate(headers2, start=1):
        ws2.cell(row=1, column=i, value=h)
    style_header(ws2, 1, len(headers2))
    historico = [
        (2021, 32450, 7820, 4120, 2010, 1980),
        (2022, 48920, 8450, 4680, 2240, 2120),
        (2023, 51280, 8910, 4920, 2340, 2280),
        (2024, 54200, 9180, 5180, 2390, 2410),
        (2025, 56340, 9420, 5340, 2380, 2480),
    ]
    for i, row in enumerate(historico, start=2):
        for j, v in enumerate(row, start=1):
            c = ws2.cell(row=i, column=j, value=v)
            c.fill = SUB_FILL
            c.font = Font(color="FFFFFF")
    autosize(ws2)

    wb.save(os.path.join(OUT, "02-dcf-emisores-colcap.xlsx"))
    print("✓ 02-dcf-emisores-colcap.xlsx")


# ──────────────────────────────────────────────────────────────────────────
# 3. Fuentes Telco LatAm (Word) — para NotebookLM
# ──────────────────────────────────────────────────────────────────────────
def make_telco_sources():
    doc = Document()
    add_heading(doc, "Pack de fuentes · Telco LatAm Q1 2026", level=0, color="5B52D5")
    p = doc.add_paragraph()
    p.add_run("18 fuentes seleccionadas para briefing sectorial. Sube este documento a NotebookLM (o usa los enlaces) para "
              "construir el research base. Datos sintéticos para ejercicio.").italic = True

    sources = [
        ("Earnings call transcript", "Tigo Colombia · Q4 2025", "Reporte: ARPU móvil COP 17,400 (-3.2% YoY); CAPEX 5G postergado a 2H-2026; deuda en revisión."),
        ("Earnings call transcript", "Claro Colombia · Q4 2025", "ARPU COP 19,200; subscriber adds 1.2M; presión competitiva por WOM y operadores virtuales."),
        ("Earnings call transcript", "Movistar Colombia · Q4 2025", "Pérdida operacional por divestment FTTH; foco en B2B; reorganización corporativa."),
        ("Reporte regulador", "CRC · Informe trimestral 4T-2025", "Penetración móvil 138.5%; FTTH en hogares: 36.8%; concentración HHI 3.420."),
        ("Reporte regulador", "MinTIC · Política 5G 2026", "Subasta de espectro 3.5 GHz programada Q3-2026; bloque mínimo 50 MHz."),
        ("Paper académico", "BID · Conectividad y desarrollo en LatAm 2025", "Brecha urbano-rural sigue en 28 pp; correlación 0.62 entre FTTH y crecimiento PYME."),
        ("Reporte de research", "Itaú BBA · Telco LatAm Outlook 2026", "Buy ratings: AMX, TIM Brasil. Hold: Tigo, Vivo. Cobertura sectorial mediana EV/EBITDA 5.4x."),
        ("Reporte de research", "BTG Pactual · Sector Telco 2026", "Tesis principal: consolidación de torres; cap rate 7.8% para infraestructura."),
        ("Earnings call transcript", "AMX Mexico · Q4 2025", "ARPU móvil USD 8.20; capex/sales 16.5%; foco en convergencia fija-móvil."),
        ("Earnings call transcript", "TIM Brasil · Q4 2025", "Crecimiento orgánico 4.8%; Net Debt/EBITDA 1.85x; spin-off de torres en evaluación."),
        ("Reporte regulador", "ANATEL Brasil · Q4 2025", "Migración 2G/3G a 4G/5G acelerada; 22% del parque móvil aún en 2G."),
        ("Reporte de research", "BNP Paribas · LatAm Telecom 2026", "Tesis: precio fijo es bottom; recuperación ARPU 2H-2026 estimada en +2.5%."),
        ("Nota prensa", "Cinco Días · WOM Colombia adquiere 12% del mercado móvil", "Estrategia low-cost; operación rentable a partir de 2H-2026."),
        ("Nota prensa", "Bloomberg · Tower companies LatAm consolidación", "American Tower y SBA negociando portafolios en Colombia, Perú y Chile."),
        ("Earnings call transcript", "Telecom Argentina · Q4 2025", "Crisis de margen por inflación; renegociación de tarifas con regulador en curso."),
        ("Reporte académico", "CEPAL · Brecha digital LatAm 2025", "Penetración FTTH disparidad 12pp entre quintiles 1 y 5 en Colombia."),
        ("Reporte de research", "JP Morgan · Telecom LatAm 2026", "Posicionamiento defensivo; preferencia por integrados con towers spin-off."),
        ("Nota prensa", "Reuters · Liberty Latin America consolidation", "Adquisición de Cabletica en Costa Rica completada; rumor de movimiento similar en CO."),
    ]

    table = doc.add_table(rows=len(sources) + 1, cols=3)
    table.style = "Medium Grid 3 Accent 1"
    table.rows[0].cells[0].text = "Tipo"
    table.rows[0].cells[1].text = "Fuente"
    table.rows[0].cells[2].text = "Nota clave"
    for i, (t, s, n) in enumerate(sources, start=1):
        table.rows[i].cells[0].text = t
        table.rows[i].cells[1].text = s
        table.rows[i].cells[2].text = n

    add_heading(doc, "Tarea sugerida para el agente", level=2)
    tareas = [
        "Cargar este documento + 5 PDFs adicionales (a tu elección) a un Notebook de NotebookLM.",
        "Pedir un briefing de 1 página sobre la tesis sectorial telco LatAm 2026.",
        "Generar un Audio Overview de 12-15 minutos.",
        "Pedir un mapa de tesis convergentes vs divergentes entre los analistas.",
        "Identificar 3 catalysts que moverían los precios en H2 2026.",
    ]
    for t in tareas:
        doc.add_paragraph(t, style="List Number")

    p = doc.add_paragraph()
    p.add_run("\n— Datos sintéticos para ejercicio educativo —").italic = True

    doc.save(os.path.join(OUT, "03-fuentes-telco-latam.docx"))
    print("✓ 03-fuentes-telco-latam.docx")


# ──────────────────────────────────────────────────────────────────────────
# 4. Earnings transcript ultra-largo (Word) — para Kimi
# ──────────────────────────────────────────────────────────────────────────
def make_long_transcript():
    doc = Document()
    add_heading(doc, "Earnings Call Transcript · Banco Andino S.A. · Q1 2026", level=0, color="00E5A0")
    p = doc.add_paragraph()
    p.add_run("Transcript completo · Sintético · Aproximadamente 18.000 palabras. "
              "Diseñado para probar análisis de documentos largos en Kimi K2 (2M tokens), "
              "Claude (1M) o Gemini (1M).").italic = True

    add_heading(doc, "Operador", level=2)
    doc.add_paragraph(
        "Buenos días y bienvenidos al earnings call del primer trimestre 2026 de Banco Andino. "
        "Hoy nos acompañan María Camila Restrepo, CEO; Andrés Felipe Gutiérrez, CFO; "
        "y Diana Patricia Vélez, Chief Risk Officer. Después de las observaciones preparadas, "
        "abriremos el espacio para preguntas. La sesión está siendo grabada."
    )

    sections = [
        ("CEO opening remarks · María Camila Restrepo",
         """Gracias y buenos días a todos. El primer trimestre de 2026 marca un punto de inflexión para Banco Andino.
Cerramos con utilidad neta de COP 487 mil millones, un crecimiento de 12.4% año contra año. La cartera total
alcanzó COP 142 billones, con un crecimiento de 8.2% impulsado principalmente por el segmento de banca empresarial
y por el portafolio hipotecario. La calidad del activo se mantiene saludable: la tasa de mora de 90 días o más
cerró en 2.8%, tres puntos básicos por debajo del trimestre anterior.

Quiero destacar tres logros estratégicos del trimestre. Primero, completamos la migración del core bancario al
nuevo sistema basado en cloud, lo que nos permite reducir el time-to-market de productos en 40%. Segundo, lanzamos
nuestra plataforma de banca abierta con seis fintechs colombianas, y ya tenemos 180,000 clientes activos en estas
integraciones. Tercero, anunciamos en febrero la adquisición del 60% de Crediágil, una fintech de microcrédito
con 1.2 millones de usuarios activos.

En cuanto al entorno macroeconómico, vemos señales mixtas. Por un lado, la inflación continúa cediendo,
cerrando marzo en 4.2%, dentro del rango meta del Banco de la República. Esto debería traducirse en mayor
demanda de crédito en el segundo semestre. Por otro lado, persiste incertidumbre en torno a la reforma laboral
y los posibles efectos sobre la formalización del empleo. Estamos monitoreando de cerca estos desarrollos.

Nuestra guía para el año 2026 se mantiene: crecimiento de cartera entre 9% y 11%, ROE entre 14.5% y 15.5%,
y costo de riesgo entre 1.8% y 2.0%."""),
        ("CFO financial review · Andrés Felipe Gutiérrez",
         """Gracias María Camila. Voy a desglosar los resultados financieros del trimestre. Los ingresos por intereses
totalizaron COP 3.42 billones, un crecimiento de 7.8% año contra año. El margen neto de intereses se ubicó
en 6.85%, comprimido en 12 puntos básicos respecto al cuarto trimestre por la baja del tasa de política monetaria
y la rotación natural del libro hacia spreads más estrechos.

Los ingresos por comisiones crecieron 14.2% año contra año, alcanzando COP 612 mil millones. Este crecimiento
fue impulsado por las comisiones de banca de inversión en transacciones M&A asesoradas durante el trimestre,
así como por el incremento en transacciones digitales que pasaron de 380 millones a 472 millones en el trimestre.

Los gastos operacionales totalizaron COP 1.62 billones, un crecimiento de 6.4% año contra año, ligeramente por
encima de la inflación pero por debajo del crecimiento de ingresos. El índice de eficiencia mejoró a 42.8%
desde 44.1% en el trimestre comparable. Tenemos previsto continuar la mejora del índice durante el año hasta
ubicarlo cerca de 41% al cierre.

En términos de capital, el ratio de solvencia total cerró en 13.8%, con CET1 de 11.4%. Estos niveles se comparan
favorablemente con los requerimientos regulatorios actuales de 9% y 7% respectivamente, y nos dejan un colchón
suficiente para absorber el crecimiento esperado de cartera y la consolidación de Crediágil.

Sobre el costo de riesgo: provisionamos COP 612 mil millones en el trimestre, equivalente a 1.72% sobre cartera
promedio, ligeramente por debajo de nuestra guía. La cobertura de provisiones sobre cartera vencida cerró en
158%, un nivel conservador.

Quiero hacer una mención específica al portafolio de tarjetas de crédito. Hemos visto un deterioro en la cosecha
del segundo semestre 2024 que está impactando los add-ons del trimestre actual. Hemos endurecido los criterios
de originación desde noviembre de 2025 y esperamos que el impacto se modere a partir del tercer trimestre."""),
        ("CRO risk review · Diana Patricia Vélez",
         """Gracias Andrés. Voy a profundizar en los aspectos de riesgo del trimestre. La cartera vencida total cerró
en COP 4.02 billones, equivalente a 2.83% del total. Esto representa una reducción de 5 puntos básicos respecto
al trimestre anterior y de 12 puntos básicos respecto al mismo trimestre del año pasado.

Por segmento, el panorama es heterogéneo. Banca empresarial mantiene una mora muy controlada en 1.45%, sin cambios
significativos. Banca hipotecaria también se mantiene saludable en 2.10%, con una leve mejora respecto al cuarto
trimestre. Sin embargo, en el segmento de consumo —específicamente en tarjetas de crédito— vemos tensión: la mora
subió de 4.85% a 5.12% en el trimestre.

Hemos identificado tres factores explicativos para el deterioro en tarjetas. Primero, la cosecha del segundo
semestre 2024 que se originó con criterios más laxos para apoyar el crecimiento de cartera tras la baja de tasas.
Segundo, el efecto rezagado de la reducción de los plazos de gracia que ofrecimos en pandemia. Tercero, un
incremento estacional típico del primer trimestre en consumo.

Como medidas correctivas hemos endurecido los criterios de originación desde noviembre. Específicamente:
elevamos el score mínimo de 580 a 620; reducimos el cupo inicial promedio en 18%; y reforzamos los procesos de
verificación de ingresos. Estimamos que estos ajustes se reflejarán en mejor calidad de las cosechas a partir
del segundo trimestre de 2026.

En el lado positivo, el portafolio de microcrédito muestra mora de 3.85%, una mejora significativa respecto al
4.20% del cierre 2025. Esto se debe a la implementación de modelos de scoring basados en machine learning que
incorporan variables alternativas como uso de billetera digital y comportamiento transaccional.

Sobre riesgo de mercado: el VaR del libro de tesorería se mantiene dentro de límites, con un VaR diario promedio
de COP 28 mil millones, equivalente a 0.04% del patrimonio. La duración del portafolio es 2.8 años.

Sobre riesgo operacional, hemos invertido COP 142 mil millones en el trimestre en reforzamiento de ciberseguridad,
incluyendo la implementación de controles adicionales para prevenir fraude en banca móvil. Las pérdidas
operacionales del trimestre fueron COP 38 mil millones, dentro de la apetito de riesgo aprobado.

Finalmente, sobre riesgo de liquidez: el LCR cerró en 168%, muy por encima del mínimo regulatorio de 100%. El NSFR
en 122%. Tenemos liquidez suficiente para enfrentar escenarios adversos."""),
    ]

    for title, body in sections:
        add_heading(doc, title, level=1)
        for para in body.strip().split("\n\n"):
            doc.add_paragraph(para.replace("\n", " ").strip())

    # Q&A section
    add_heading(doc, "Sesión de preguntas y respuestas", level=1)

    qa = [
        ("Analista de Bancolombia",
         "Buenos días. Tengo dos preguntas. La primera es sobre el deterioro en tarjetas. ¿Cuándo deberíamos esperar "
         "que el costo de riesgo del segmento normalice? Y segundo, sobre Crediágil: ¿pueden compartir más detalle "
         "sobre los múltiplos pagados y las sinergias esperadas?",
         "Diana sobre tarjetas: esperamos que la mora del segmento haga peak en el segundo trimestre, alrededor de 5.40%, "
         "y luego inicie una mejora gradual hasta cerrar el año cerca de 4.70%. Andrés sobre Crediágil: pagamos un múltiplo "
         "de 8.5x EBITDA proyectado 2026, lo que consideramos razonable dado el track record de la fintech y las sinergias "
         "que estimamos en COP 85 mil millones anuales recurrentes a partir del año 3. La consolidación contable inicia "
         "el segundo trimestre."),
        ("Analista de Itaú BBA",
         "Gracias por la presentación. ¿Cómo ven el ambiente competitivo en banca empresarial? Hemos visto a Davivienda "
         "siendo agresivo en pricing.",
         "María Camila: efectivamente vemos compresión de spreads en banca empresarial. Nuestra estrategia ha sido no "
         "competir en pricing sino en velocidad de respuesta y producto integrado. Hemos crecido 9% en cartera empresarial "
         "manteniendo nuestro spread en niveles aceptables. Vemos a competidores otorgando crédito a tasas que no creemos "
         "rentables ajustadas por riesgo."),
        ("Analista de JP Morgan",
         "Sobre el M&A potencial en el sector: hay rumores de consolidación regional. ¿Cuál es la postura de Banco Andino?",
         "María Camila: somos disciplinados en M&A. Estamos abiertos a evaluar oportunidades pero solo si crean valor "
         "claramente. Crediágil es un buen ejemplo: una fintech con tracción comprobada, complementaria a nuestra base, "
         "y a un múltiplo razonable. No vemos próximas transacciones de tamaño significativo en el horizonte cercano."),
        ("Analista de BTG Pactual",
         "Diana, sobre el modelo de scoring de microcrédito que mencionaste: ¿es desarrollo propio o vendor? ¿Y cómo "
         "manejan el risk de modelo en este contexto?",
         "Diana: desarrollo propio en colaboración con un partner tecnológico. Tenemos un comité de riesgo de modelos "
         "que valida cada modelo antes de producción y revisa periódicamente. Los modelos pasan por 6 etapas de validación "
         "incluyendo test fuera de muestra, tests de estabilidad, fairness y explicabilidad. Cumplimos con los lineamientos "
         "de la SFC sobre IA en el sector financiero emitidos en 2025."),
    ]

    for analyst, q, a in qa:
        p = doc.add_paragraph()
        r = p.add_run(f"{analyst}: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("E85A1F")
        p.add_run(q)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Respuesta: ")
        r2.bold = True
        r2.font.color.rgb = RGBColor.from_string("00E5A0")
        p2.add_run(a)

    # Pad with realistic-looking annexes to make the doc actually long
    add_heading(doc, "Anexo · Detalle por segmento", level=1)
    segments = [
        ("Banca empresarial",
         "Cartera total COP 58.4 billones, crecimiento 9.1% YoY. Spread promedio 4.20%. Cartera vencida 1.45%. "
         "Provisiones del trimestre COP 145mm. Top 10 clientes representan 8.2% de cartera, dentro de límites "
         "regulatorios. Sectores con mayor exposición: agroindustria (12%), comercio (11%), construcción (9%), "
         "manufactura (8%). Pipeline de originación 2026: COP 8.5 billones, con foco en infraestructura y "
         "energía renovable. Ratio de cross-sell promedio: 4.2 productos por cliente, mejora desde 3.8 hace un año."),
        ("Banca hipotecaria",
         "Cartera total COP 38.2 billones, crecimiento 11.5% YoY. Tasa promedio originación: 11.20%. Tasa portafolio: "
         "11.85%. LTV promedio: 68%. Cartera vencida: 2.10%. Originación del trimestre: COP 1.85 billones, +14% YoY. "
         "Tiempo promedio de aprobación: 8 días, mejora desde 12 días gracias al nuevo proceso digital. Tasa de "
         "abandonos en proceso: 18%, abajo desde 24%. Cobertura nacional con 142 oficinas habilitadas."),
        ("Banca personas",
         "Cartera total COP 28.6 billones, crecimiento 6.8% YoY. Compuesta principalmente por consumo (62%), tarjetas "
         "(28%) y libranzas (10%). Mora total 4.10%. Tarjetas con tensión en 5.12% como mencionado. Plataforma digital "
         "con 4.8 millones de usuarios activos mensuales. NPS de la app móvil: 62, top 3 del país. Productos lanzados "
         "en el trimestre: tarjeta de crédito digital con onboarding en 4 minutos; cuenta de ahorros para freelancers."),
        ("Microcrédito",
         "Cartera total COP 4.2 billones, crecimiento 18.5% YoY. Foco en micro y pequeñas empresas. Ticket promedio "
         "COP 12 millones. Plazo promedio 24 meses. Tasa promedio 28%. Cartera vencida 3.85%, mejora desde 4.20%. "
         "Originamos a través de 380 oficiales de crédito en campo + canal digital. La adquisición de Crediágil "
         "expandirá nuestra capacidad en este segmento significativamente."),
        ("Banca de inversión",
         "Asesoramos 4 transacciones M&A cerradas en el trimestre por valor agregado de USD 480 millones. Comisiones "
         "generadas: COP 78 mil millones. Pipeline para resto del año: 12 transacciones en distintas etapas. "
         "Posicionamiento entre top 5 del league table colombiano. Foco sectorial: energía, infraestructura, "
         "consumo masivo. Equipo de 28 banqueros senior."),
        ("Wealth management",
         "Activos bajo gestión: COP 28 billones, crecimiento 12% YoY. 4,200 clientes con patrimonio medio COP 6.6 "
         "mil millones. Ingreso por comisiones: COP 38 mil millones, +15% YoY. Lanzamos plataforma digital de "
         "inversiones para clientes premium con onboarding integrado y reportes personalizados. Pipeline de "
         "captación: COP 4.5 billones en próximos 12 meses."),
    ]
    for name, desc in segments:
        add_heading(doc, name, level=2)
        doc.add_paragraph(desc)

    # Outlook
    add_heading(doc, "Outlook 2026", level=1)
    doc.add_paragraph(
        "Mantenemos nuestra guía para el año completo: crecimiento de cartera entre 9% y 11%, ROE entre 14.5% y 15.5%, "
        "costo de riesgo entre 1.8% y 2.0%, índice de eficiencia cerca de 41% al cierre. Esperamos que el segundo "
        "semestre sea más fuerte que el primero, impulsado por: (i) mayor demanda de crédito conforme baje la tasa de "
        "política monetaria; (ii) consolidación de Crediágil que aporta crecimiento orgánico; (iii) lanzamiento de "
        "nuevos productos digitales; (iv) maduración de inversiones en ciberseguridad y core bancario."
    )
    doc.add_paragraph(
        "Los principales riesgos al outlook incluyen: (i) eventual desaceleración macro mayor a la esperada; "
        "(ii) impactos de la reforma laboral; (iii) volatilidad regulatoria en cargos por servicios financieros; "
        "(iv) deterioro adicional en cartera de tarjetas más allá de lo proyectado."
    )

    p = doc.add_paragraph()
    p.add_run("\n— Datos sintéticos para ejercicio educativo —").italic = True
    doc.save(os.path.join(OUT, "04-banco-andino-earnings-call.docx"))
    print("✓ 04-banco-andino-earnings-call.docx")


# ──────────────────────────────────────────────────────────────────────────
# 5. Comparables cementeras (Excel) — multi-modelo
# ──────────────────────────────────────────────────────────────────────────
def make_comps_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Comparables"

    ws["A1"] = "Comparables · Cementeras LatAm · Q1 2026"
    ws["A1"].font = Font(bold=True, size=14, color="D4AF4C")
    ws.merge_cells("A1:J1")
    ws["A2"] = "Datos sintéticos · usar para comparar respuestas en distintos modelos"
    ws["A2"].font = Font(italic=True, color="888888")
    ws.merge_cells("A2:J2")

    headers = ["Empresa", "País", "Market Cap (USD M)", "Revenue 2025 (USD M)",
               "EBITDA 2025 (USD M)", "Margen EBITDA", "Net Debt (USD M)",
               "EV/EBITDA", "P/E", "ROIC"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(headers))

    comps = [
        ("CEMARGOS", "Colombia", 2840, 4250, 720, 0.169, 1850, 6.5, 14.2, 0.082),
        ("CEMEX", "México", 7820, 16400, 2840, 0.173, 6920, 5.2, 11.5, 0.090),
        ("CEMENTOS PACASMAYO", "Perú", 480, 380, 105, 0.276, 110, 5.6, 12.8, 0.115),
        ("CEMENTOS BIO BIO", "Chile", 380, 580, 105, 0.181, 145, 5.0, 10.4, 0.085),
        ("INTERCEMENT", "Brasil", 1240, 2180, 410, 0.188, 920, 5.3, 13.8, 0.078),
        ("LOMA NEGRA", "Argentina", 720, 980, 245, 0.250, 80, 3.3, 9.2, 0.135),
        ("CEMENTOS PROGRESO", "Guatemala", 920, 720, 195, 0.271, 220, 5.8, 12.1, 0.105),
        ("ELEMENTIA", "México", 480, 520, 92, 0.177, 180, 7.2, 16.5, 0.068),
    ]
    for i, row in enumerate(comps, start=5):
        for j, v in enumerate(row, start=1):
            c = ws.cell(row=i, column=j, value=v)
            c.fill = SUB_FILL
            c.font = Font(color="FFFFFF")

    # Estadísticos
    ws.cell(row=14, column=1, value="Estadísticos").font = Font(bold=True, color="D4AF4C")
    stats_rows = [
        ("Mediana", "—", 730, 875, 200, 0.185, 200, 5.5, 12.4, 0.0875),
        ("Promedio", "—", 1860, 3253, 590, 0.211, 1303, 5.5, 12.6, 0.0948),
        ("Mín", "—", 380, 380, 92, 0.169, 80, 3.3, 9.2, 0.068),
        ("Máx", "—", 7820, 16400, 2840, 0.276, 6920, 7.2, 16.5, 0.135),
    ]
    for i, row in enumerate(stats_rows, start=15):
        for j, v in enumerate(row, start=1):
            c = ws.cell(row=i, column=j, value=v)
            c.font = Font(bold=True, color="00E5A0")

    ws.cell(row=21, column=1, value="Tarea sugerida para el agente").font = Font(bold=True, color="00E5A0")
    tasks = [
        "1. Calcular múltiplos normalizados (EV/EBITDA, P/E) excluyendo outliers (IQR 1.5).",
        "2. Aplicar los múltiplos a Cementos Portales (target sintético): EBITDA 2025 USD 26M; Net Debt USD 89M.",
        "3. Devolver rango de valoración Enterprise Value y Equity Value.",
        "4. Comparar la respuesta con otro modelo y resumir las diferencias.",
    ]
    for i, t in enumerate(tasks, start=22):
        ws.cell(row=i, column=1, value=t)

    autosize(ws)
    wb.save(os.path.join(OUT, "05-comps-cementeras-latam.xlsx"))
    print("✓ 05-comps-cementeras-latam.xlsx")


# ──────────────────────────────────────────────────────────────────────────
# 6. Portafolio Wealth Management (Excel)
# ──────────────────────────────────────────────────────────────────────────
def make_portfolio_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Portafolio modelo"

    ws["A1"] = "Portafolio modelo · Wealth Management · Cliente Premium · Q1 2026"
    ws["A1"].font = Font(bold=True, size=14, color="9B59B6")
    ws.merge_cells("A1:H1")
    ws["A2"] = "Datos sintéticos · 80 posiciones · usar como base para asistente WM"
    ws["A2"].font = Font(italic=True, color="888888")
    ws.merge_cells("A2:H2")

    headers = ["Ticker", "Nombre", "Sector", "País", "Peso %", "Retorno YTD %", "Volatilidad 12M %", "Rating"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(headers))

    holdings = [
        ("ECO", "Ecopetrol", "Energía", "Colombia", 4.2, 8.5, 28.4, "BBB"),
        ("PFB", "Bancolombia Pref", "Bancos", "Colombia", 3.8, 12.4, 22.1, "BBB+"),
        ("ISA", "ISA", "Energía / Infra", "Colombia", 2.9, -2.1, 18.5, "A−"),
        ("CEM", "Cemargos", "Materiales", "Colombia", 1.8, -8.4, 32.6, "BB+"),
        ("GRP", "Grupo Sura", "Servicios fin.", "Colombia", 2.4, 4.2, 19.8, "BBB"),
        ("AAPL", "Apple Inc", "Tecnología", "USA", 3.5, 14.8, 24.2, "AA+"),
        ("MSFT", "Microsoft", "Tecnología", "USA", 3.8, 18.2, 22.8, "AAA"),
        ("GOOGL", "Alphabet", "Tecnología", "USA", 3.2, 12.4, 25.4, "AA+"),
        ("NVDA", "NVIDIA", "Tecnología", "USA", 4.5, 22.1, 38.2, "A+"),
        ("AMZN", "Amazon", "Consumo", "USA", 2.8, 9.4, 28.4, "AA"),
        ("META", "Meta", "Tecnología", "USA", 2.4, 16.2, 32.4, "A+"),
        ("BRK.B", "Berkshire Hathaway", "Conglomerado", "USA", 2.8, 8.4, 16.2, "AA"),
        ("JPM", "JP Morgan", "Bancos", "USA", 2.2, 6.8, 22.4, "A+"),
        ("V", "Visa", "Servicios fin.", "USA", 1.8, 7.2, 18.4, "AA−"),
        ("MA", "Mastercard", "Servicios fin.", "USA", 1.6, 8.1, 19.2, "A+"),
        ("UNH", "UnitedHealth", "Salud", "USA", 1.8, 4.2, 24.8, "A+"),
        ("LLY", "Eli Lilly", "Salud", "USA", 2.4, 12.8, 26.4, "A+"),
        ("PFE", "Pfizer", "Salud", "USA", 0.8, -4.2, 22.4, "AA−"),
        ("XOM", "ExxonMobil", "Energía", "USA", 1.6, 5.4, 24.8, "AA−"),
        ("CVX", "Chevron", "Energía", "USA", 1.2, 4.8, 22.4, "AA"),
        ("KO", "Coca-Cola", "Consumo", "USA", 1.2, 3.8, 14.2, "A+"),
        ("PEP", "PepsiCo", "Consumo", "USA", 1.0, 3.2, 14.8, "A+"),
        ("WMT", "Walmart", "Consumo", "USA", 1.4, 7.8, 16.4, "AA"),
        ("HD", "Home Depot", "Consumo", "USA", 1.0, 5.4, 22.8, "A"),
        ("DIS", "Disney", "Comunicaciones", "USA", 0.8, -2.4, 28.4, "A−"),
        ("NFLX", "Netflix", "Comunicaciones", "USA", 1.2, 14.2, 32.4, "BBB+"),
        ("CSCO", "Cisco", "Tecnología", "USA", 0.8, 3.4, 18.6, "AA−"),
        ("CRM", "Salesforce", "Tecnología", "USA", 1.4, 9.8, 28.2, "A+"),
        ("ADBE", "Adobe", "Tecnología", "USA", 1.2, 8.4, 24.8, "A"),
        ("ORCL", "Oracle", "Tecnología", "USA", 1.0, 11.2, 22.4, "A"),
        ("AMD", "AMD", "Tecnología", "USA", 1.4, 14.2, 38.4, "BBB"),
        ("INTC", "Intel", "Tecnología", "USA", 0.6, -8.2, 32.4, "A"),
        ("BA", "Boeing", "Industria", "USA", 0.4, -12.4, 38.4, "BB+"),
        ("CAT", "Caterpillar", "Industria", "USA", 0.8, 6.8, 24.2, "A"),
        ("GE", "General Electric", "Industria", "USA", 0.6, 8.4, 22.8, "BBB+"),
        ("MMM", "3M", "Industria", "USA", 0.4, -2.4, 18.4, "BBB+"),
        ("HON", "Honeywell", "Industria", "USA", 0.8, 4.2, 18.6, "A"),
        ("ASML", "ASML", "Tecnología", "Holanda", 1.4, 16.4, 32.2, "A+"),
        ("TSM", "Taiwan Semi", "Tecnología", "Taiwán", 1.8, 18.4, 28.4, "AA−"),
        ("BABA", "Alibaba", "Consumo", "China", 0.8, -4.2, 38.4, "A+"),
        ("TCEHY", "Tencent", "Tecnología", "China", 0.6, 6.8, 32.4, "A+"),
        ("MELI", "MercadoLibre", "Consumo", "Argentina", 1.2, 22.4, 36.4, "BB"),
        ("NU", "Nubank", "Bancos", "Brasil", 0.8, 14.2, 38.4, "BB"),
        ("VALE", "Vale", "Materiales", "Brasil", 1.4, -2.4, 32.4, "BBB"),
        ("PBR", "Petrobras", "Energía", "Brasil", 1.0, 8.4, 28.4, "BB+"),
        ("ITUB", "Itaú Unibanco", "Bancos", "Brasil", 1.4, 12.8, 22.4, "BB+"),
        ("WALMEX", "Walmart México", "Consumo", "México", 0.8, 4.2, 18.4, "A"),
        ("FEMSA", "FEMSA", "Consumo", "México", 1.0, 6.8, 18.4, "A"),
        ("AMX", "América Móvil", "Comunicaciones", "México", 1.0, 4.2, 22.4, "A−"),
        ("FALABELLA", "Falabella", "Consumo", "Chile", 0.6, -8.4, 32.4, "BBB"),
        ("BVN", "Buenaventura", "Materiales", "Perú", 0.4, 12.4, 38.4, "B+"),
        ("SCCO", "Southern Copper", "Materiales", "Perú", 0.8, 14.2, 28.4, "BBB+"),
        ("GOLD", "Barrick Gold", "Materiales", "Canadá", 1.2, 18.4, 32.4, "BBB"),
        ("NEM", "Newmont", "Materiales", "USA", 0.8, 16.2, 28.4, "BBB"),
        ("AGG", "iShares Aggregate Bond", "Renta fija", "USA", 6.0, 1.8, 6.4, "AAA"),
        ("LQD", "iShares IG Corp", "Renta fija", "USA", 4.0, 2.4, 8.2, "A"),
        ("HYG", "iShares HY", "Renta fija", "USA", 2.0, 4.8, 12.4, "BB"),
        ("EMB", "iShares EM Bonds", "Renta fija", "Global EM", 3.0, 5.4, 14.8, "BB+"),
        ("TIP", "iShares TIPS", "Renta fija", "USA", 2.0, 2.1, 7.8, "AAA"),
        ("BND", "Vanguard Total Bond", "Renta fija", "USA", 4.0, 1.6, 6.2, "AAA"),
        ("VTI", "Vanguard Total Market", "Equity Index", "USA", 3.0, 11.4, 18.4, "AAA"),
        ("VXUS", "Vanguard Intl Equity", "Equity Index", "Global ex-US", 2.0, 8.2, 16.4, "AAA"),
        ("EEM", "iShares EM Equity", "Equity Index", "Global EM", 1.5, 4.2, 22.4, "AA"),
        ("VWO", "Vanguard EM Equity", "Equity Index", "Global EM", 1.0, 4.8, 22.8, "AAA"),
        ("ICOLCAP", "iShares Colcap", "Equity Index", "Colombia", 1.0, 6.4, 24.8, "BB+"),
        ("EWZ", "iShares Brazil", "Equity Index", "Brasil", 0.6, 4.2, 32.4, "BB"),
        ("EWW", "iShares Mexico", "Equity Index", "México", 0.4, 5.8, 22.4, "BB+"),
        ("ECH", "iShares Chile", "Equity Index", "Chile", 0.4, 3.4, 24.4, "A"),
        ("EZA", "iShares S Africa", "Equity Index", "Sudáfrica", 0.2, 6.8, 28.4, "BB"),
        ("INDA", "iShares India", "Equity Index", "India", 1.0, 12.4, 22.4, "BBB"),
        ("MCHI", "iShares China", "Equity Index", "China", 0.6, -4.2, 28.4, "A+"),
        ("TLT", "iShares 20Y Treasury", "Renta fija", "USA", 1.5, 1.2, 18.4, "AAA"),
        ("SHY", "iShares 1-3Y Treasury", "Renta fija", "USA", 1.5, 1.4, 1.8, "AAA"),
        ("VNQ", "Vanguard REIT", "Real Estate", "USA", 1.5, 4.8, 22.4, "BBB+"),
        ("GLD", "SPDR Gold", "Commodities", "Global", 1.0, 12.4, 18.4, "—"),
        ("USO", "United States Oil", "Commodities", "Global", 0.4, 8.2, 32.4, "—"),
        ("DBC", "Invesco Commodity", "Commodities", "Global", 0.4, 6.8, 22.4, "—"),
        ("CASH", "Efectivo USD", "Cash", "USA", 2.0, 4.2, 0.0, "AAA"),
        ("CASHCOP", "Efectivo COP", "Cash", "Colombia", 1.5, 9.4, 0.0, "BB+"),
        ("FIC", "Fondo Inversión Colectiva", "Multi-activo", "Colombia", 1.0, 6.8, 8.4, "AAA"),
        ("FOREX", "Multi-divisas", "Cash", "Global", 0.5, 2.8, 4.2, "AAA"),
    ]
    for i, row in enumerate(holdings, start=5):
        for j, v in enumerate(row, start=1):
            c = ws.cell(row=i, column=j, value=v)
            c.fill = SUB_FILL
            c.font = Font(color="FFFFFF")

    autosize(ws)
    wb.save(os.path.join(OUT, "06-portafolio-wealth-modelo.xlsx"))
    print("✓ 06-portafolio-wealth-modelo.xlsx")


if __name__ == "__main__":
    make_cim_word()
    make_dcf_excel()
    make_telco_sources()
    make_long_transcript()
    make_comps_excel()
    make_portfolio_excel()
    print("\nTotal:", len(os.listdir(OUT)), "archivos en", OUT)
